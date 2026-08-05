from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


OUT = Path('/Users/patrick/workspace/variableway/factory/fire-skills/chinese-font-and-poetry-research.docx')

BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
INK = '0B2545'
MUTED = '5F6B76'
LIGHT_BLUE = 'E8EEF5'
LIGHT_GRAY = 'F2F4F7'
GOLD = 'A46A1F'
WHITE = 'FFFFFF'


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths)))
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '120')
    tblInd.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def set_run_font(run, name='Arial Unicode MS', size=None, color=None, bold=None, italic=None, east_asia='Arial Unicode MS'):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), east_asia)
    rFonts.set(qn('w:hint'), 'eastAsia')
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    lang.set(qn('w:val'), 'en-US')
    lang.set(qn('w:eastAsia'), 'zh-CN')
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_font(style, name='Arial Unicode MS', size=11, color='000000', bold=False, east_asia='Arial Unicode MS'):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), east_asia)
    rFonts.set(qn('w:hint'), 'eastAsia')
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    lang.set(qn('w:val'), 'en-US')
    lang.set(qn('w:eastAsia'), 'zh-CN')


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True, size=10.5):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), color)
    rPr.append(color_el)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    rFonts.set(qn('w:hint'), 'eastAsia')
    rPr.append(rFonts)
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'en-US')
    lang.set(qn('w:eastAsia'), 'zh-CN')
    rPr.append(lang)
    new_run.append(rPr)
    text_el = OxmlElement('w:t')
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run_font(run, size=9, color=MUTED)


def add_para(doc, text='', style=None, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=11)
    return p


def add_label_para(doc, label, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(label)
    set_run_font(r, size=11, color=INK, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=11)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Inches(0.375 if level == 0 else 0.625)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size={1:16,2:13,3:12}[level], color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(label + '  ')
    set_run_font(r, size=10.5, color=GOLD, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_repo_table(doc):
    rows = [
        ('1', 'w3c/clreq', '799', '中文排版标准：横竖排、标点、行首行尾规则和 Web/电子书支持。', 'https://github.com/w3c/clreq'),
        ('2', 'skishore/makemeahanzi', '2.6k', '9000+ 汉字的部件拆分、字源类型、笔顺和 SVG 路径。', 'https://github.com/skishore/makemeahanzi'),
        ('3', 'SCUT-DLVCLab/MCCD', '36', '约 33 万书法字图，标注 10 种书体、15 个朝代和 142 位书法家。', 'https://github.com/SCUT-DLVCLab/MCCD'),
        ('4', 'adobe-fonts/source-han-sans', '16.9k', '思源黑体；现代黑体、字重、可变字体和中日韩字形差异。', 'https://github.com/adobe-fonts/source-han-sans'),
        ('5', 'adobe-fonts/source-han-serif', '9.6k', '思源宋体；宋体/明体、印刷字体和地区字形研究。', 'https://github.com/adobe-fonts/source-han-serif'),
        ('6', 'lxgw/LxgwWenKai', '25.2k', '霞鹜文楷；现代楷体与屏幕阅读，适合诗词和人文内容。', 'https://github.com/lxgw/LxgwWenKai'),
        ('7', 'atelier-anchor/smiley-sans', '14.6k', '得意黑；现代几何、展示型中文黑体。', 'https://github.com/atelier-anchor/smiley-sans'),
        ('8', 'TakWolf/ark-pixel-font', '4.6k', '方舟像素字体；研究中文字体进入游戏和像素媒介后的变化。', 'https://github.com/TakWolf/ark-pixel-font'),
        ('9', 'jaywcjlove/free-font', '4.0k', '字体目录，按黑体、宋体、楷体、艺术体、手绘体等分类。', 'https://github.com/jaywcjlove/free-font'),
        ('10', 'chinese-poetry/chinese-poetry', '52.7k', '唐诗、宋诗、宋词等大型诗词语料，适合排印、检索和数字人文。', 'https://github.com/chinese-poetry/chinese-poetry'),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [500, 2500, 900, 5460])
    headers = ['#', '仓库', '约 Star', '研究价值']
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, LIGHT_BLUE)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0,2) else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        set_run_font(r, size=10, color=INK, bold=True)
    set_repeat_table_header(table.rows[0])
    for rank, repo, stars, value, url in rows:
        cells = table.add_row().cells
        for i, txt in enumerate([rank, repo, stars, value]):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if i == 1:
                add_hyperlink(p, txt, url, size=9.5)
            else:
                r = p.add_run(txt)
                set_run_font(r, size=9.5, color=INK)
            if i in (0,2):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def add_modern_table(doc):
    rows = [
        ('诗云 / Poetry Cloud', 'Cohenjikan/shiyun', '把近百万首诗放进可漫游的三维星系，展示朝代、诗人、格律和赠答网络。', '空间叙事、关系网络、计算诗学'),
        ('她写过 / She Wrote', 'teenycao/shewrote', '给诗人增加性别与人物身份信息，讨论女性诗人如何在文献中被遗漏。', '数字人文、数据偏差、性别史'),
        ('AsPoem', 'meetqy/aspoem', '用拼音、注释、译文、赏析、纠错和打印功能重新设计诗词阅读。', '现代阅读体验、教育'),
        ('COPE', 'LingDong-/cope', '将平仄、押韵、格律检查与风格比较做成诗歌创作 IDE。', '创作工具、传统规则的现代使用'),
        ('京墨', 'hefengbao/jingmo', '把诗词与汉字、节气、传统色、节日和人物放到同一个文化阅读应用中。', '文化生活方式、跨知识连接'),
        ('Poetry Skill', 'Wscats/poetry-skills', '将大型诗词数据接入自然语言查询、比较、分析和创作辅助。', 'AI 入口、语义检索'),
        ('诗泉 API', 'palemoky/chinese-poetry-api', '通过 REST/GraphQL 提供全文搜索、简繁切换、分类和随机诗词服务。', '基础设施、可复用 API'),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1700, 1900, 3600, 2160])
    headers = ['项目', '仓库', '做法', '现代意义']
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, LIGHT_BLUE)
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=9.5, color=INK, bold=True)
    set_repeat_table_header(table.rows[0])
    urls = {
        'Cohenjikan/shiyun':'https://github.com/Cohenjikan/shiyun',
        'teenycao/shewrote':'https://github.com/teenycao/shewrote',
        'meetqy/aspoem':'https://github.com/meetqy/aspoem',
        'LingDong-/cope':'https://github.com/LingDong-/cope',
        'hefengbao/jingmo':'https://github.com/hefengbao/jingmo',
        'Wscats/poetry-skills':'https://github.com/Wscats/poetry-skills',
        'palemoky/chinese-poetry-api':'https://github.com/palemoky/chinese-poetry-api',
    }
    for name, repo, approach, meaning in rows:
        cells = table.add_row().cells
        for i, txt in enumerate([name, repo, approach, meaning]):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if i == 1:
                add_hyperlink(p, txt, urls[repo], size=9)
            else:
                r = p.add_run(txt)
                set_run_font(r, size=9)
    return table


def add_module_table(doc):
    rows = [
        ('当代情境搜索', '用户输入“搬家、失眠、异地、告别、重新开始”等生活处境，检索相关真实诗句。'),
        ('诗人时空地图', '展示作品写作地点、行旅路径、人物关系和今天对应的真实地点。'),
        ('被遗漏者专题', '延续 She Wrote 的思路，增加女性、僧人、地方诗人和无名作者等视角。'),
        ('中文字体实验室', '用宋体、楷体、仿宋、黑体、手写体、像素体和竖排呈现同一首诗，解释字体选择。'),
        ('当代回应层', '允许读者附上照片、声音、城市、日期或一句现代回应，形成新的注释层。'),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2200, 7160])
    for i, h in enumerate(['模块', '功能']):
        c = table.rows[0].cells[i]
        set_cell_shading(c, LIGHT_BLUE)
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=10, color=INK, bold=True)
    set_repeat_table_header(table.rows[0])
    for name, desc in rows:
        cells = table.add_row().cells
        for i, txt in enumerate([name, desc]):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(txt)
            set_run_font(r, size=10, color=INK, bold=(i == 0))
    return table


def configure_styles(doc):
    normal = doc.styles['Normal']
    style_font(normal, size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        st = doc.styles[f'Heading {level}']
        style_font(st, size=size, color=color, bold=True)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15
        st.paragraph_format.keep_with_next = True
    for name in ['List Bullet', 'List Bullet 2', 'List Number']:
        st = doc.styles[name]
        style_font(st, size=11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    # Quiet running furniture for the multi-page reference report.
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    hr = header.add_run('中文字体研究 · chinese-poetry 现代化方向')
    set_run_font(hr, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run('Research note  ·  ')
    set_run_font(fr, size=9, color=MUTED)
    add_page_field(footer)

    # Cover / editorial opening.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(100)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('研究整理')
    set_run_font(r, size=11, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('中文字体研究与\nchinese-poetry 现代化方向')
    set_run_font(r, name='Arial Unicode MS', size=27, color=INK, bold=True, east_asia='Arial Unicode MS')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run('从字体史、字形数据与古典诗词，到当代阅读和数字人文项目')
    set_run_font(r, size=13, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [1800, 7560])
    meta_rows = [
        ('文档目的', '整理本次关于中文字体研究、GitHub 资源和 chinese-poetry 现代化意义的完整讨论'),
        ('数据核验', 'GitHub 仓库信息主要核验于 2026-07-22；Star 数会持续变化'),
        ('研究范围', '字体历史与分类、汉字结构、书法数据、中文排版、诗词语料、现代化产品方向'),
        ('核心判断', '价值不在于再做一个搜索页，而在于建立“诗词 ↔ 当代经验”的解释与参与层'),
    ]
    for row, (label, value) in zip(meta.rows, meta_rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(label)
        set_run_font(r0, size=10, color=INK, bold=True)
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(value)
        set_run_font(r1, size=10)

    doc.add_page_break()

    add_heading(doc, '一、讨论结论', 1)
    add_para(doc, '如果想做一个中文字体研究，GitHub 上没有一个官方的“中文字体史 Top 10”，但可以把资源分成四层：排版规范、字形与书法数据、代表性字体、诗词语料。真正有研究价值的项目，应把这四层放在一起，而不是只做字体下载或诗词检索。')
    add_note(doc, '核心建议', '以 chinese-poetry 作为诗词语料底座，以字体与排版作为视觉研究层，再建立“古典文本如何回应今天的生活经验”的解释层。')
    add_para(doc, 'GitHub 的优势是源码、数据、字体和可复现实验；它的不足是缺少完整的中文字体史叙事。因此，历史部分仍需要结合书法史、印刷史、字体设计史和正式出版物。')

    add_heading(doc, '二、中文字体研究的 GitHub Top 10（研究取向）', 1)
    add_para(doc, '下面不是纯粹按 Star 排名，而是按“历史与理论价值、字体类型覆盖、数据可用性、活跃度和对项目的帮助”综合筛选。')
    add_repo_table(doc)
    add_para(doc, '说明：MCCD 的 Star 数不高，但它对书法史与历史字体研究的价值很高；它包含约 33 万书法字符图、10 种书体、15 个朝代和 142 位书法家。', after=4)
    add_para(doc, '版权提醒：free-font 收录的每一款字体有自己的授权条款；MCCD 限制为非商业研究；Make Me a Hanzi 的字典数据和字形数据也来自不同来源，使用前应分别核对许可证。', after=8)

    add_heading(doc, '三、建议的字体研究路径', 1)
    add_number(doc, '先读 W3C clreq（https://github.com/w3c/clreq），建立中文横排、竖排、标点、行间和电子书排版的基本框架。')
    add_number(doc, '用 Make Me a Hanzi（https://github.com/skishore/makemeahanzi）研究汉字的部件、笔顺、字源类型和 SVG 字形结构。')
    add_number(doc, '用 MCCD 观察书体、朝代和书法家的视觉差异，形成“篆、隶、楷、行、草”等历史样本。')
    add_number(doc, '用思源黑体、思源宋体、霞鹜文楷、得意黑、方舟像素体建立现代字体对照组。')
    add_number(doc, '最后将同一首诗放入不同字体、不同字号、不同字距、竖排和横排中，记录阅读感受与视觉差异。')
    add_para(doc, '可重点观察：宋体的印刷感、楷体的书写感、黑体的现代公共性、仿宋的文献感、手写体的亲密感，以及像素体对“屏幕时代”的回应。')

    add_heading(doc, '四、chinese-poetry 的现有生态与现代转译', 1)
    add_para(doc, 'chinese-poetry 本身提供了约 5.5 万首唐诗、26 万首宋诗、2.1 万首宋词以及其他古典文集。它最初的现代意义是让庞大的古典文集以 JSON 形式可复制、可检索、可构建应用。现在更值得推进的是：在原始数据之上建立新的解释、关系和参与方式。')
    add_modern_table(doc)

    add_heading(doc, '五、什么才算“更多现代意义”', 1)
    add_para(doc, '现代意义不只是把古诗词放进一个更漂亮的页面，也不只是让大模型自动生成一首诗。它至少包括以下几个方向：')
    add_bullet(doc, '从“背诵名句”转向“解释当代处境”：让用户通过搬家、失眠、异地、告别、重新开始等真实情境进入诗词。')
    add_bullet(doc, '从“诗人名录”转向“谁被记录、谁被遗漏”：增加性别、社会身份、地域、阶层和文献保存偏差。')
    add_bullet(doc, '从“平面列表”转向“时间与空间”：把诗人的行旅、赠答、地理和历史事件连接起来。')
    add_bullet(doc, '从“固定文本”转向“可参与的注释”：允许读者添加今天的照片、声音、地点和回应，但保留出处和版本。')
    add_bullet(doc, '从“字体只是皮肤”转向“字体是解释”：让字体、排版方向、字距和书体成为诗意的一部分。')

    add_heading(doc, '六、最值得做的新项目：此刻有诗 / Poetry for Now', 1)
    add_para(doc, '建议项目定位为“古典诗词与当代生活经验的语义接口”。用户不再主要按作者、朝代和标题搜索，而是输入自己正在经历的事情，系统从真实诗词中找到相关文本，再解释它为什么与今天产生联系。')
    add_heading(doc, '一个典型使用场景', 2)
    add_para(doc, '用户输入：“我要离开生活了十年的城市，既期待又舍不得。”系统返回若干真实诗句，并分别给出：原文、作者与出处、写作背景、关键意象、与现代处境的相似处和不同处，以及一段明确标注为 AI 辅助的现代解读。')
    add_heading(doc, '建议功能模块', 2)
    add_module_table(doc)
    add_heading(doc, '为什么这个方向有价值', 2)
    add_bullet(doc, '它把古典诗词从“文化库存”变成可以参与当代生活的语言资源。')
    add_bullet(doc, '它保留真实出处，避免把 AI 生成内容伪装成古人作品。')
    add_bullet(doc, '它可以自然结合中文字体、排版、书法史和数字人文，不会局限在单一技术上。')
    add_bullet(doc, '它既可以做公共文化产品，也可以做字体研究、教育研究和文学研究的实验平台。')

    add_heading(doc, '七、数据架构建议', 1)
    add_para(doc, '不要直接改写 chinese-poetry 的原始数据；建议在其上增加可追溯的 enrichment layer：')
    arch = doc.add_table(rows=1, cols=2)
    set_table_geometry(arch, [2200, 7160])
    for i, h in enumerate(['数据层', '内容']):
        c = arch.rows[0].cells[i]
        set_cell_shading(c, LIGHT_BLUE)
        r = c.paragraphs[0].add_run(h)
        set_run_font(r, size=10, color=INK, bold=True)
    set_repeat_table_header(arch.rows[0])
    arch_rows = [
        ('原始层', '保留 chinese-poetry 原始 JSON、文件路径、版本、来源和校订记录。'),
        ('身份层', '作者别名、生卒年、朝代、籍贯、性别（有证据才标注）和人物关系。'),
        ('语义层', '意象、主题、情绪、生活情境、修辞、体裁、格律和相关诗句。'),
        ('视觉层', '字体、书体、字形版本、横竖排、字号、字距、颜色和版式实验。'),
        ('参与层', '读者注释、现代回应、纠错、出处、可信度和版本历史。'),
    ]
    for name, desc in arch_rows:
        cells = arch.add_row().cells
        for i, txt in enumerate([name, desc]):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(txt)
            set_run_font(r, size=10, color=INK, bold=(i == 0))
    add_note(doc, '数据治理', '中国古典诗词文本大多已进入公版范围，但数据库整理、现代注释、现当代诗歌文本和字体文件仍可能有独立版权。每条衍生数据都应保留 provenance、许可证和置信度。')

    add_heading(doc, '八、MVP 路线图', 1)
    add_heading(doc, '阶段 1：可信语料', 2)
    add_bullet(doc, '固定一个版本的 chinese-poetry，记录 commit、文件和字段。')
    add_bullet(doc, '建立作者、作品、朝代、出处和纠错表。')
    add_bullet(doc, '先做 1000 首高质量样本，不急于覆盖全部数据。')
    add_heading(doc, '阶段 2：阅读与检索', 2)
    add_bullet(doc, '实现“生活情境 + 意象 + 主题 + 诗句”的混合检索。')
    add_bullet(doc, '提供原文、注释、译文、出处和现代解释的分层显示。')
    add_bullet(doc, '加入宋体、楷体、黑体和竖排等字体实验。')
    add_heading(doc, '阶段 3：数字人文与参与', 2)
    add_bullet(doc, '加入地图、时间线、诗人关系网络和被遗漏者专题。')
    add_bullet(doc, '让用户提交现代回应，并建立审核、引用和版本机制。')
    add_bullet(doc, '将数据层开放成 API，让其他教育、字体和文化项目复用。')

    add_heading(doc, '九、结论', 1)
    add_para(doc, '最有辨识度的方向，是把“中文字体研究”与“古典诗词现代化”放在同一个体验里：研究字形如何改变诗的阅读方式，也研究古典诗词如何继续解释今天的人。')
    add_para(doc, '一句话概括：不要只把诗词保存下来，而要让它在新的字体、新的屏幕、新的生活处境和新的读者身上继续发生。', after=12)

    add_heading(doc, '附录：主要链接', 1)
    links = [
        ('中文诗词数据库', 'https://github.com/chinese-poetry/chinese-poetry'),
        ('W3C 中文排版要求', 'https://github.com/w3c/clreq'),
        ('Make Me a Hanzi', 'https://github.com/skishore/makemeahanzi'),
        ('MCCD 书法历史数据集', 'https://github.com/SCUT-DLVCLab/MCCD'),
        ('思源黑体', 'https://github.com/adobe-fonts/source-han-sans'),
        ('思源宋体', 'https://github.com/adobe-fonts/source-han-serif'),
        ('霞鹜文楷', 'https://github.com/lxgw/LxgwWenKai'),
        ('诗云', 'https://github.com/Cohenjikan/shiyun'),
        ('她写过', 'https://github.com/teenycao/shewrote'),
        ('AsPoem', 'https://github.com/meetqy/aspoem'),
        ('COPE', 'https://github.com/LingDong-/cope'),
        ('京墨', 'https://github.com/hefengbao/jingmo'),
        ('诗泉 API', 'https://github.com/palemoky/chinese-poetry-api'),
    ]
    for label, url in links:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        add_hyperlink(p, label, url, size=10)
        r = p.add_run('  ' + url)
        set_run_font(r, size=9, color=MUTED)

    # Core properties / metadata.
    doc.core_properties.title = '中文字体研究与 chinese-poetry 现代化方向'
    doc.core_properties.subject = '中文字体、古典诗词、GitHub 资源与现代化项目方向'
    doc.core_properties.author = 'Codex'
    doc.core_properties.keywords = '中文字体, chinese-poetry, 诗词, 数字人文, 中文排版'
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
